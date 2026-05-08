"""Text extraction from uploaded files (NFR-09, NFR-12).

Dispatches to the right library based on file extension:

    .txt   → plain UTF-8 decode
    .docx  → python-docx  (paragraph text, joined by newlines)
    .pdf   → pdfplumber   (page text, joined by newlines; None pages skipped)

Raises
------
fastapi.HTTPException(400)
    Unsupported file type (NFR-08 / NFR-09) or extraction failure.

Design notes
------------
* Only the file *extension* is used for dispatch; MIME sniffing is left to
  the router layer (NFR-09) so this module stays a pure preprocessing
  concern (NFR-12).
* All libraries operate on ``io.BytesIO`` — no temp files written to disk
  (NFR-10: no persistent storage of user content).
* pdf pages that yield ``None`` from pdfplumber (e.g. image-only pages) are
  silently skipped; the caller receives whatever text was extractable.
"""

from __future__ import annotations

import io
import logging

from fastapi import HTTPException, status

logger = logging.getLogger(__name__)

# File extensions this module can handle.
SUPPORTED_EXTENSIONS: frozenset[str] = frozenset({".txt", ".docx", ".pdf"})


def extract_text(content: bytes, filename: str) -> str:
    """Return the plain-text content of *content* using the right extractor.

    Parameters
    ----------
    content:
        Raw file bytes, already read and size-validated by the caller.
    filename:
        Original filename (used only to determine the extension).

    Returns
    -------
    str
        Extracted plain text, ready for tokenisation by the spell-check engine.

    Raises
    ------
    HTTPException(400)
        If the extension is not supported, or if the library raises an
        exception during extraction (e.g. corrupt file).
    """
    ext = _extension(filename)

    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=(
                f"Unsupported file type '{ext}'. "
                f"Accepted formats: {', '.join(sorted(SUPPORTED_EXTENSIONS))}."
            ),
        )

    try:
        if ext == ".txt":
            return _extract_txt(content)
        if ext == ".docx":
            return _extract_docx(content)
        if ext == ".pdf":
            return _extract_pdf(content)
    except HTTPException:
        raise
    except Exception as exc:
        logger.exception("Extraction failed for '%s': %s", filename, exc)
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Could not extract text from '{filename}': {exc}",
        ) from exc

    # Unreachable — satisfies mypy's exhaustiveness check.
    raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Unknown file type.")


# ---------------------------------------------------------------------------
# Private extractors
# ---------------------------------------------------------------------------

def _extension(filename: str) -> str:
    """Return the lowercased file extension including the dot, e.g. ``'.pdf'``."""
    dot = filename.rfind(".")
    if dot == -1:
        return ""
    return filename[dot:].lower()


def _extract_txt(content: bytes) -> str:
    """Decode raw bytes as UTF-8, replacing undecodable bytes (NFR-10)."""
    return content.decode("utf-8", errors="replace")


def _extract_docx(content: bytes) -> str:
    """Extract paragraph text from a .docx file using python-docx.

    Only top-level ``Document.paragraphs`` are collected. Tables and
    text-boxes are out of scope for v1.0 but can be added here without
    touching the router or engine layers.
    """
    import docx  # python-docx; imported lazily so the module loads even if
                 # the package is absent in test environments that mock it.

    doc = docx.Document(io.BytesIO(content))
    lines = [para.text for para in doc.paragraphs]
    return "\n".join(lines)


def _extract_pdf(content: bytes) -> str:
    """Extract text from every page of a PDF using pdfplumber.

    Pages that return ``None`` from ``extract_text()`` (e.g. scanned /
    image-only pages) are skipped rather than raising; the caller receives
    whatever text was extractable and the spell checker runs on that subset.
    """
    import pdfplumber  # lazy import — same rationale as _extract_docx.

    pages: list[str] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)

    return "\n".join(pages)
